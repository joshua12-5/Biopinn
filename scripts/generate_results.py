#!/usr/bin/env python3
"""CLI: generate the full manuscript Results & Discussion asset pack, in APA
(7th edition) style.

Loads the trained checkpoint (+ optionally a w_phys=0 ablation baseline) and
the held-out test set, computes every figure and table via src/results.py,
and writes them to results/paper/:
  - fig_*.png / .pdf   (300 DPI, 12 figures; no caption baked into the image --
    APA figures carry data/axes/legend only, see FIGURE_CAPTIONS_<mode>.md)
  - table_*.csv        (11 tables)
  - BIOPINN_results_tables_<mode>.docx   (tables compiled under Research
    Question headings -- see RQ_OUTLINE below -- APA three-line borders,
    "Table N" + italicized title, "Note." footnotes)
  - FIGURE_CAPTIONS_<mode>.md            (APA figure captions: "Figure N" +
    italicized title, external to the image files themselves)
  - results_manifest_<mode>.json         (source/config/key-values per asset)

--numbering controls the caption/filename scheme (<mode> above) for the
original Fig/Table 4.X set:
  - "chapter"    (default): dissertation-chapter-style, e.g. Table 4.1,
    fig_4_1_concentration_heatmap.png -- identical to this script's original
    numbering.
  - "sequential": plain APA journal-article style, e.g. Table 1,
    fig_1_concentration_heatmap.png. Figures and tables are numbered in
    independent sequences (Figure 1..10, Table 1..9), per APA convention.
Run the script twice, once per --numbering value, to get both -- filenames
never collide between the two modes, so both can live in results/paper/
at once. The two RQ1b/RQ1c assets (added to directly answer the sub-parts of
Research Question 1 that the original Fig/Table 4.X set didn't cover as a
function of nanoparticle diameter) keep a fixed "rq1b"/"rq1c" label under
both numbering modes, since they aren't part of that chapter sequence.

The compiled .docx groups every table (and references every figure) under
four Research Question headings, matching the project's Statement of the
Problem -- see RQ_OUTLINE below for exactly what evidence answers each
question and sub-question.

Never retrains -- consumes artifacts/ exactly as the other scripts/*.py do.
Table 8 (ablation) is skipped with a clear note if no baseline checkpoint
is present (run scripts/run_ablation.py first to produce one).

Usage:
    python scripts/generate_results.py [--experiment NAME] [--n-jobs N] [--numbering {chapter,sequential}]
"""

from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

import numpy as np
import torch
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from src import results as R
from src.biology import evaluate_h4_hypothesis, predict_concentration_field, viability_map
from src.config import load_config, resolve_path
from src.evaluate import evaluate_h2_hypothesis, resolve_test_simulations
from src.microenvironment import radial_grid
from src.model import load_checkpoint
from src.optimize import optimize_all_radii, speedup_study

APA_FONT = "Times New Roman"

# Compiled-document outline: groups every table (and references every
# figure) under the project's four Statement-of-the-Problem research
# questions, instead of the old flat Fig/Table 4.1..4.10 sequence. Each
# entry is one of:
#   ("h1", text)                    -- top-level "Research Question N" heading
#   ("h2", text)                    -- sub-question heading (RQ1's a/b/c)
#   ("table", chapter_number)       -- one table via _add_table_to_doc
#   ("figref", chapter_number)      -- a one-line pointer to the standalone
#                                      figure file (figures aren't embedded
#                                      in the compiled document -- see
#                                      FIGURE_CAPTIONS_<mode>.md, matching
#                                      this project's existing APA
#                                      figures-are-external-files design)
#   ("text", key)                   -- a short prose paragraph, built from
#                                      real computed values (see _RQ4_TEXT_KEY)
RQ_OUTLINE = [
    ("h1", "Research Question 1: What is the diffusion coefficient effect of "
           "simulated nanoparticle diameter on:"),
    ("h2", "a. Drug penetration depth"),
    ("table", "4.2"), ("table", "4.4"), ("figref", "4.6"),
    ("h2", "b. Concentration distribution across tumor zones"),
    ("table", "rq1b"), ("figref", "rq1b"),
    ("h2", "c. Sub-therapeutic regions"),
    ("table", "rq1c"), ("figref", "rq1c"),
    ("table", "4.6"), ("figref", "4.9"),
    ("h1", "Research Question 2: Accuracy of the trained BIOPINN model in "
           "reproducing the finite-difference reference concentration field "
           "on the held-out test set (RMSE, MAE, R², L2 relative error)"),
    ("table", "4.3"), ("figref", "4.4"), ("figref", "4.3"),
    ("h1", "Research Question 3: Predicted cell survivability and cytotoxicity "
           "during drug penetration and exposure"),
    ("table", "4.5"), ("figref", "4.7"), ("figref", "4.8"),
    ("h1", "Research Question 4: What combination of nanoparticle diameter and "
           "dosing produces the greatest predicted tumor cell kill, and why?"),
    ("table", "4.7"), ("figref", "4.10"), ("text", "rq4_interpretation"),
    ("h1", "Supporting Material"),
    ("table", "4.1"), ("table", "4.8"), ("table", "4.9"),
    ("figref", "4.1"), ("figref", "4.2"), ("figref", "4.5"),
]


def _rq4_interpretation(grid_result: dict, config: dict) -> str:
    """Short, computed (not canned) interpretation of RQ4's optimum: where it
    sits within the searched grid, and the Stokes-Einstein reasoning (D_free
    scales as 1/d_NP -- already established in src/microenvironment.py and
    tested in tests/test_microenvironment.py) for why that side of the range
    tends to win."""
    opt_cfg = config["optimization"]
    d_lo, d_hi = opt_cfg["d_NP_grid_nm"]
    c_lo, c_hi = opt_cfg["C0_grid_uM"]
    d_star, c_star, eta = grid_result["d_NP_star_nm"], grid_result["C0_star_uM"], grid_result["max_eta"]

    d_frac = (d_star - d_lo) / (d_hi - d_lo) if d_hi > d_lo else 0.5
    c_frac = (c_star - c_lo) / (c_hi - c_lo) if c_hi > c_lo else 0.5
    d_side = "smaller" if d_frac < 0.5 else "larger"
    diffusion_word = "faster" if d_side == "smaller" else "slower"
    reach_word = "more" if d_side == "smaller" else "less"
    c_side = "lower" if c_frac < 0.5 else "higher"

    return (
        f"Across the searched grid (d_NP in [{d_lo:g}, {d_hi:g}] nm, C0 in "
        f"[{c_lo:g}, {c_hi:g}] μM) at the baseline tumor radius, the combination "
        f"maximizing the volume-averaged kill fraction η was d_NP* = {d_star:.1f} nm "
        f"and C0* = {c_star:.2f} μM, reaching η = {eta * 100.0:.1f}%. By the "
        f"Stokes-Einstein relation (Table RQ1b/RQ1c and Table 4.2), free diffusivity "
        f"D_free scales as 1/d_NP, so on diffusion alone a {d_side} nanoparticle would "
        f"be expected to diffuse {diffusion_word} and reach {reach_word} of the tumor "
        f"volume within the exposure window; the optimal diameter found here sits "
        f"toward that {d_side} end of the searched range, and the optimal dose toward "
        f"the {c_side} end of the searched concentration range. Kill fraction also "
        f"depends on local exposure duration and decay rate, not diffusion reach alone, "
        f"so this diffusion-based reasoning explains the general tendency rather than "
        f"fully determining the exact optimum."
    )


def _sanitize_for_json(obj):
    """Recursively converts numpy/pandas scalars to native Python types (same
    conversion _default performs for json.dump) and replaces NaN/Infinity
    with None -- json.dump's default allow_nan=True would otherwise emit a
    literal NaN/Infinity token, which is not valid JSON and breaks strict
    parsers (e.g. a div-by-zero guard in table_4_6's percentage-difference
    column, or an empty resistance zone elsewhere, both legitimate results
    at some parameter combinations)."""
    if isinstance(obj, dict):
        return {k: _sanitize_for_json(v) for k, v in obj.items()}
    if isinstance(obj, (list, tuple)):
        return [_sanitize_for_json(v) for v in obj]
    if isinstance(obj, float):
        return obj if np.isfinite(obj) else None
    if hasattr(obj, "tolist"):
        return _sanitize_for_json(obj.tolist())
    if hasattr(obj, "item"):
        return _sanitize_for_json(obj.item())
    return obj


def _seq(chapter_number: str) -> str:
    """'4.3' -> '3'. The chapter numbering (Fig/Table 4.1..4.N) is already
    sequential within the chapter with no gaps, so the plain APA sequential
    number is just this suffix -- no separate remapping table needed."""
    return chapter_number.split(".")[1]


def _display_number(chapter_number: str, numbering: str) -> str:
    """The number shown in a caption/heading: 'Table 4.3' vs. 'Table 3'."""
    return chapter_number if numbering == "chapter" else _seq(chapter_number)


def _file_number(chapter_number: str, numbering: str) -> str:
    """The number embedded in a filename: 'fig_4_3_...' vs. 'fig_3_...'."""
    return chapter_number.replace(".", "_") if numbering == "chapter" else _seq(chapter_number)


def _set_cell_border(cell, **edges) -> None:
    """Low-level docx border control (python-docx has no high-level API for
    per-edge table borders). Each kwarg is one of top/bottom/left/right,
    valued with a dict of OOXML border attributes, e.g. {"val": "single",
    "sz": 8, "color": "000000"} or {"val": "nil"} to remove that edge."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge, attrs in edges.items():
        edge_el = tcBorders.find(qn(f"w:{edge}"))
        if edge_el is None:
            edge_el = OxmlElement(f"w:{edge}")
            tcBorders.append(edge_el)
        for key, value in attrs.items():
            edge_el.set(qn(f"w:{key}"), str(value))


def _apply_apa_three_line_borders(table) -> None:
    """APA 7 "three-line table" style (Publication Manual Table 7.7): no
    vertical rules anywhere, and horizontal rules only above the header row,
    below the header row, and below the final data row -- nothing else."""
    NIL = {"val": "nil"}
    RULE = {"val": "single", "sz": 8, "color": "000000"}
    n_rows = len(table.rows)
    for r_idx, row in enumerate(table.rows):
        for cell in row.cells:
            _set_cell_border(
                cell,
                left=NIL,
                right=NIL,
                top=RULE if r_idx == 0 else NIL,
                bottom=RULE if r_idx in (0, n_rows - 1) else NIL,
            )


def _add_table_to_doc(doc: Document, display_number: str, caption: str, df, footnote: str | None = None) -> None:
    """Write one APA-style table: bold 'Table N' on its own line, an
    italicized title below it, a borderless three-line table (see
    _apply_apa_three_line_borders), and an optional 'Note. ...' footnote
    with only the 'Note.' label italicized, per APA 7 table conventions."""
    number_para = doc.add_paragraph()
    number_run = number_para.add_run(f"Table {display_number}")
    number_run.bold = True
    number_run.font.name = APA_FONT
    number_run.font.size = Pt(12)

    title_para = doc.add_paragraph()
    title_run = title_para.add_run(caption)
    title_run.italic = True
    title_run.font.name = APA_FONT
    title_run.font.size = Pt(12)

    table = doc.add_table(rows=1, cols=len(df.columns))
    header_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        header_cells[i].text = str(col)
        header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in header_cells[i].paragraphs[0].runs:
            run.font.name = APA_FONT
            run.font.size = Pt(11)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = f"{value:.4g}" if isinstance(value, float) else str(value)
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cells[i].paragraphs[0].runs:
                run.font.name = APA_FONT
                run.font.size = Pt(11)
    _apply_apa_three_line_borders(table)

    if footnote:
        note = doc.add_paragraph()
        label_run = note.add_run("Note. ")
        label_run.italic = True
        label_run.font.name = APA_FONT
        label_run.font.size = Pt(10)
        body_run = note.add_run(footnote)
        body_run.font.name = APA_FONT
        body_run.font.size = Pt(10)
    doc.add_paragraph()


def _write_figure_captions(manifest_figures: dict, figure_order: list[str], numbering: str, output_dir: Path) -> Path:
    """APA-style figure caption list (bold 'Figure N', italicized title)
    as a standalone markdown file -- the figures themselves stay plain
    PNG/PDF files with no caption baked in, so this is where the caption
    text actually lives."""
    lines = ["# BIOPINN Figure Captions (APA style)", ""]
    for chapter_number in figure_order:
        entry = manifest_figures[chapter_number]
        png_name = Path(entry["files"]["png"]).name
        pdf_name = Path(entry["files"]["pdf"]).name
        lines.append(f"**Figure {entry['display_number']}**")
        lines.append("")
        lines.append(f"*{entry.get('caption', '')}*")
        lines.append("")
        lines.append(f"File: `{png_name}` / `{pdf_name}`")
        lines.append("")
    path = output_dir / f"FIGURE_CAPTIONS_{numbering}.md"
    path.write_text("\n".join(lines), encoding="utf-8")
    return path


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--experiment", default=None, help="Experiment config name (default: default_config.yaml)")
    parser.add_argument("--n-jobs", type=int, default=1, help="Parallel workers for re-solving the test set's FDM reference")
    parser.add_argument(
        "--numbering",
        choices=("chapter", "sequential"),
        default="chapter",
        help="Caption/filename numbering scheme: 'chapter' (Table 4.1, default, unchanged from before) "
        "or 'sequential' (Table 1, plain APA journal-article style). Run once per value to get both.",
    )
    args = parser.parse_args()

    config = load_config(args.experiment)
    paper_cfg = config["paper"]

    checkpoint_path = resolve_path(config, "model_checkpoint")
    if not checkpoint_path.exists():
        print(f"No checkpoint found at {checkpoint_path}.")
        print("Run notebooks/biopinn_train.ipynb on Colab first, then drop its artifacts here.")
        sys.exit(1)

    processed_dir = resolve_path(config, "processed")
    if not (processed_dir / "test.npz").exists():
        print(f"No processed test set found at {processed_dir}.")
        print("Run notebooks/biopinn_train.ipynb (or src.data_pipeline.build_dataset) first.")
        sys.exit(1)

    print(f"Loading checkpoint: {checkpoint_path}")
    model = load_checkpoint(str(checkpoint_path), config)
    with open(resolve_path(config, "normalization_stats"), encoding="utf-8") as f:
        norm_stats = json.load(f)

    baseline_checkpoint_path = checkpoint_path.with_name(checkpoint_path.stem + "_baseline" + checkpoint_path.suffix)
    baseline_model = load_checkpoint(str(baseline_checkpoint_path), config) if baseline_checkpoint_path.exists() else None
    if baseline_model is None:
        print(f"No ablation baseline checkpoint at {baseline_checkpoint_path} -- Table 4.8 will be skipped.")
        print("Run scripts/run_ablation.py first to enable it.")

    print(f"Re-solving FDM reference for the test set (n_jobs={args.n_jobs})...")
    sims = resolve_test_simulations(config, n_jobs=args.n_jobs)
    print(f"  {len(sims)} test simulations resolved.")

    with np.load(processed_dir / "test.npz") as test_npz:
        collocation_X = torch.as_tensor(test_npz["collocation_X"], dtype=torch.float32)

    output_dir = resolve_path(config, "results") / "paper"
    output_dir.mkdir(parents=True, exist_ok=True)
    manifest: dict = {"config_used": args.experiment or "default_config.yaml", "baseline_params": R.baseline_params(config), "figures": {}, "tables": {}}

    # ----------------------------------------------------------------- #
    # Figures
    # ----------------------------------------------------------------- #
    print("\nComputing figures...")

    print("  Fig 4.1 -- spatiotemporal concentration heatmap")
    fig, meta = R.fig_4_1_concentration_heatmap(model, config, norm_stats)
    stem = f"fig_{_file_number('4.1', args.numbering)}_concentration_heatmap"
    manifest["figures"]["4.1"] = {"files": R._save_figure(fig, output_dir, stem), "display_number": _display_number("4.1", args.numbering), **meta}

    print("  Fig 4.2 -- radial concentration profiles")
    fig, meta = R.fig_4_2_radial_profiles(model, config, norm_stats)
    stem = f"fig_{_file_number('4.2', args.numbering)}_radial_concentration_profiles"
    manifest["figures"]["4.2"] = {"files": R._save_figure(fig, output_dir, stem), "display_number": _display_number("4.2", args.numbering), **meta}

    print("  solving the nanoparticle-diameter FDM sweep (shared by Table 4.1/4.2/4.4 and Fig 4.3)...")
    diameter_sweep = R.solve_diameter_sweep(config)

    print("  Fig 4.3 -- PINN vs. FDM comparison at t=24hr")
    fig, meta = R.fig_4_3_pinn_vs_fdm_t24(model, config, norm_stats, diameter_sweep=diameter_sweep)
    stem = f"fig_{_file_number('4.3', args.numbering)}_pinn_vs_fdm_profile_comparison"
    manifest["figures"]["4.3"] = {"files": R._save_figure(fig, output_dir, stem), "display_number": _display_number("4.3", args.numbering), **meta}

    print("  Fig 4.4 -- predicted vs. reference scatter (full test set)")
    fig, meta = R.fig_4_4_scatter_pred_vs_ref(model, config, norm_stats, sims)
    stem = f"fig_{_file_number('4.4', args.numbering)}_predicted_vs_reference_scatter"
    manifest["figures"]["4.4"] = {"files": R._save_figure(fig, output_dir, stem), "display_number": _display_number("4.4", args.numbering), **meta}

    print("  Fig 4.5 -- training loss convergence")
    history = R.load_training_history(config)
    fig, meta = R.fig_4_5_training_loss(history, config)
    stem = f"fig_{_file_number('4.5', args.numbering)}_training_loss_convergence"
    manifest["figures"]["4.5"] = {"files": R._save_figure(fig, output_dir, stem), "display_number": _display_number("4.5", args.numbering), **meta}

    print("  Fig 4.6 -- penetration depth vs. time")
    fig, meta = R.fig_4_6_penetration_vs_time(model, config, norm_stats)
    stem = f"fig_{_file_number('4.6', args.numbering)}_penetration_depth_vs_time"
    manifest["figures"]["4.6"] = {"files": R._save_figure(fig, output_dir, stem), "display_number": _display_number("4.6", args.numbering), **meta}

    print("  Fig 4.7 -- spatial viability at t=72hr")
    fig, meta = R.fig_4_7_viability_t72(model, config, norm_stats)
    stem = f"fig_{_file_number('4.7', args.numbering)}_spatial_viability"
    manifest["figures"]["4.7"] = {"files": R._save_figure(fig, output_dir, stem), "display_number": _display_number("4.7", args.numbering), **meta}

    print("  Fig 4.8 -- cytotoxicity evolution")
    fig, meta = R.fig_4_8_cytotoxicity_evolution(model, config, norm_stats)
    stem = f"fig_{_file_number('4.8', args.numbering)}_cytotoxicity_evolution"
    manifest["figures"]["4.8"] = {"files": R._save_figure(fig, output_dir, stem), "display_number": _display_number("4.8", args.numbering), **meta}

    print("  Fig RQ1b -- concentration by tumor zone across nanoparticle diameters")
    fig, meta = R.fig_rq1b_concentration_by_zone(config, diameter_sweep=diameter_sweep)
    stem = "fig_rq1b_concentration_by_zone"
    manifest["figures"]["rq1b"] = {"files": R._save_figure(fig, output_dir, stem), "display_number": "RQ1b", **meta}

    print("  Fig RQ1c -- sub-therapeutic region vs. nanoparticle diameter")
    fig, meta = R.fig_rq1c_subtherapeutic_vs_diameter(config, diameter_sweep=diameter_sweep)
    stem = "fig_rq1c_subtherapeutic_vs_diameter"
    manifest["figures"]["rq1c"] = {"files": R._save_figure(fig, output_dir, stem), "display_number": "RQ1c", **meta}

    print("  solving the heterogeneous/homogeneous D_eff comparison (shared by Table 4.6 and Fig 4.9)...")
    hetero_homog = R.solve_hetero_and_homogeneous(config)

    print("  Fig 4.9 -- heterogeneous vs. homogeneous D_eff")
    fig, meta = R.fig_4_9_hetero_vs_homog(config, solved=hetero_homog)
    stem = f"fig_{_file_number('4.9', args.numbering)}_heterogeneous_vs_homogeneous"
    manifest["figures"]["4.9"] = {"files": R._save_figure(fig, output_dir, stem), "display_number": _display_number("4.9", args.numbering), **meta}

    print("  running the grid search at the baseline radius (shared by Table 4.7 and Fig 4.10)...")
    radius_results = optimize_all_radii(model, config, norm_stats)
    baseline_R = R.baseline_params(config)["R_um"]

    print("  Fig 4.10 -- treatment effectiveness surface")
    fig, meta = R.fig_4_10_effectiveness_surface(model, config, norm_stats, grid_result=radius_results[baseline_R])
    stem = f"fig_{_file_number('4.10', args.numbering)}_effectiveness_surface"
    manifest["figures"]["4.10"] = {"files": R._save_figure(fig, output_dir, stem), "display_number": _display_number("4.10", args.numbering), **meta}

    # ----------------------------------------------------------------- #
    # Tables
    # ----------------------------------------------------------------- #
    print("\nComputing tables...")
    tables: dict[str, tuple] = {}

    print("  Table 4.1 -- FDM simulation summary")
    tables["4.1"] = R.table_4_1_fdm_summary(config, diameter_sweep=diameter_sweep)

    print("  Table 4.2 -- max penetration depth vs. diameter")
    tables["4.2"] = R.table_4_2_penetration_vs_diameter(config, diameter_sweep=diameter_sweep)

    print("  Table 4.3 -- PINN test-set metrics")
    tables["4.3"] = R.table_4_3_pinn_metrics(model, config, norm_stats, sims, collocation_X)

    print("  Table 4.4 -- penetration depth analysis")
    tables["4.4"] = R.table_4_4_penetration_analysis(model, config, norm_stats, diameter_sweep=diameter_sweep)

    print("  Table 4.5 -- viability summary")
    tables["4.5"] = R.table_4_5_viability_summary(model, config, norm_stats)

    print("  Table RQ1b -- concentration by tumor zone across nanoparticle diameters")
    tables["rq1b"] = R.table_rq1b_concentration_by_zone(config, diameter_sweep=diameter_sweep)

    print("  Table RQ1c -- sub-therapeutic region vs. nanoparticle diameter")
    tables["rq1c"] = R.table_rq1c_subtherapeutic_by_diameter(config, diameter_sweep=diameter_sweep)

    print("  Table 4.6 -- heterogeneous vs. homogeneous comparison")
    tables["4.6"] = R.table_4_6_hetero_vs_homog(config, solved=hetero_homog)

    print("  Table 4.7 -- optimization results")
    tables["4.7"] = R.table_4_7_optimization_results(model, config, norm_stats, radius_results=radius_results)

    if baseline_model is not None:
        print("  Table 4.8 -- ablation study")
        tables["4.8"] = R.table_4_8_ablation(model, baseline_model, config, norm_stats, sims, collocation_X)
    else:
        tables["4.8"] = None

    for number, result in tables.items():
        if result is None:
            continue
        df, meta = result
        # rq1b/rq1c aren't part of the Fig/Table 4.X chapter sequence -- they're
        # tied to specific SOP sub-questions, so they keep a fixed "RQ1b"/"RQ1c"
        # label under both numbering modes instead of feeding _file_number/
        # _display_number (which assume the "N.M" chapter-number format).
        if number.startswith("rq"):
            csv_path = output_dir / f"table_{number}.csv"
            display_number = "RQ" + number[2:]
        else:
            csv_path = output_dir / f"table_{_file_number(number, args.numbering)}.csv"
            display_number = _display_number(number, args.numbering)
        df.to_csv(csv_path, index=False, encoding="utf-8")
        manifest["tables"][number] = {"csv": str(csv_path), "display_number": display_number, **meta}

    # ----------------------------------------------------------------- #
    # Hypothesis summary (Table 4.9) -- assembled from the tables/figures above
    # ----------------------------------------------------------------- #
    print("  Table 4.9 -- hypothesis evaluation summary")
    h1_meta = tables["4.3"][1]
    h2 = evaluate_h2_hypothesis(model, norm_stats, config)

    baseline_p = R.baseline_params(config)
    r_h4 = radial_grid(baseline_p["R_um"], 200, config["fdm"]["r_min_um"])
    t_h4 = np.linspace(0.0, baseline_p["t_max_hr"], 60)
    C_h4 = predict_concentration_field(model, r_h4, t_h4, baseline_p, norm_stats)
    V_h4 = viability_map(C_h4, t_h4, config)
    h4 = evaluate_h4_hypothesis(V_h4, r_h4, t_h4, config, R_um=baseline_p["R_um"])

    h3_meta = tables["4.6"][1]

    if tables["4.8"] is not None:
        h5_pass = tables["4.8"][1]["H5_pass"]
    else:
        h5_pass = False

    print(f"  running the PINN-vs-FDM speedup study ({paper_cfg['hypotheses']['H6_speedup_n_combinations']} combinations)...")
    speedup = speedup_study(model, config, norm_stats, n_combinations=paper_cfg["hypotheses"]["H6_speedup_n_combinations"])
    table_4_7_times = [row["Computation Time (sec)"] for row in tables["4.7"][0].to_dict("records")]
    h6_pass = bool(
        speedup["speedup_ratio_mean"] >= 10 ** config["evaluation"]["hypotheses"]["H6_speedup_orders_of_magnitude"]
        and float(np.mean(table_4_7_times)) <= paper_cfg["hypotheses"]["H6_max_pinn_time_s"]
    )
    manifest["H6_speedup_study"] = speedup

    hypotheses = {
        "H1": {
            "name": "PINN Accuracy", "expected": "R² > 0.990, RMSE < 0.05 μM (target)",
            "evidence": "Table 4.3", "pass": bool(h1_meta["H1_pass_on_pooled_metrics"]),
        },
        "H2": {
            "name": "NP Size vs. Depth", "expected": "72hr penetration-depth difference (10nm vs. 200nm) > 100 μm",
            "evidence": "Table 4.4", "pass": bool(h2["pass"]),
        },
        "H3": {
            "name": "Spatial Heterogeneity", "expected": f"Steeper gradient, sub-therapeutic zone >= {h3_meta['H3_threshold_pct']:.0f}% larger",
            "evidence": "Table 4.6", "pass": bool(h3_meta["H3_pass"]),
        },
        "H4": {
            "name": "Viability Pattern", "expected": "Rim < 20%, Core > 60% at t=72hr",
            "evidence": "Table 4.5", "pass": bool(h4["overall_pass"]),
        },
        "H5": {
            "name": "Physics vs. Data-Driven", "expected": ">= 10x lower PDE residual, Wilcoxon significant",
            "evidence": "Table 4.8", "pass": bool(h5_pass),
        },
        "H6": {
            "name": "Surrogate Speedup", "expected": ">= 2 orders of magnitude speedup, <= 5s per configuration",
            "evidence": "Table 4.7", "pass": h6_pass,
        },
    }
    table_4_9 = R.table_4_9_hypothesis_summary(hypotheses)
    table_4_9_csv_path = output_dir / f"table_{_file_number('4.9', args.numbering)}.csv"
    table_4_9.to_csv(table_4_9_csv_path, index=False, encoding="utf-8")
    manifest["tables"]["4.9"] = {"csv": str(table_4_9_csv_path), "display_number": _display_number("4.9", args.numbering), "hypotheses": hypotheses}

    # ----------------------------------------------------------------- #
    # Compiled Word document (APA-style tables: bold "Table N", italicized
    # title, borderless three-line table, "Note." footnotes -- see
    # _add_table_to_doc)
    # ----------------------------------------------------------------- #
    docx_path = output_dir / f"BIOPINN_results_tables_{args.numbering}.docx"
    print(f"\nCompiling {docx_path.name}...")
    doc = Document()
    doc.add_heading("BIOPINN — Results & Discussion", level=1)
    doc.add_paragraph(
        f"Every value in this document was computed from the trained checkpoint, the held-out "
        f"{len(sims)}-simulation test set, and the analysis routines in src/ -- see "
        f"results_manifest_{args.numbering}.json in this folder for the exact source and "
        f"configuration behind every number. Tables and figures are organized below by the "
        f"research question they answer; see FIGURE_CAPTIONS_{args.numbering}.md for the figures "
        f"themselves (not embedded here, per APA figure convention)."
    )

    table_captions = {
        "4.1": f"FDM Simulation Summary Statistics (Baseline: R={baseline_p['R_um']:.0f} μm, d_NP={baseline_p['d_NP_nm']:.0f} nm)",
        "4.2": "Maximum Penetration Depth at t=72hr Across Nanoparticle Diameters",
        "4.3": f"PINN Test Set Evaluation Metrics (Mean ± SD, N={len(sims)} test simulations)",
        "4.4": f"Penetration Depth Analysis Across Nanoparticle Sizes (R={baseline_p['R_um']:.0f}μm, C0={baseline_p['C0_uM']:.0f}μM, t={baseline_p['t_max_hr']:.0f}hr)",
        "4.5": f"Viability Analysis Summary at t={baseline_p['t_max_hr']:.0f}hr (R={baseline_p['R_um']:.0f}μm, C0={baseline_p['C0_uM']:.0f}μM)",
        "4.6": f"Heterogeneous vs. Homogeneous D_eff Model Comparison (d_NP={baseline_p['d_NP_nm']:.0f}nm, R={baseline_p['R_um']:.0f}μm, t={baseline_p['t_max_hr']:.0f}hr)",
        "4.7": "Optimization Results: Optimal Nanoparticle Parameters for Different Tumor Sizes",
        "4.8": "Ablation Study: PINN vs. Unconstrained NN Baseline",
        "4.9": "Hypothesis Evaluation Summary",
        "rq1b": f"Mean Concentration by Tumor Zone Across Nanoparticle Diameters (R={baseline_p['R_um']:.0f}μm, C0={baseline_p['C0_uM']:.0f}μM, t={baseline_p['t_max_hr']:.0f}hr)",
        "rq1c": f"Sub-therapeutic Tumor Region vs. Nanoparticle Diameter (R={baseline_p['R_um']:.0f}μm, C0={baseline_p['C0_uM']:.0f}μM, t={baseline_p['t_max_hr']:.0f}hr)",
    }
    rq4_text = _rq4_interpretation(radius_results[baseline_R], config)

    for block in RQ_OUTLINE:
        kind = block[0]
        if kind == "h1":
            doc.add_heading(block[1], level=2)
            continue
        if kind == "h2":
            doc.add_heading(block[1], level=3)
            continue
        if kind == "text":
            doc.add_paragraph(rq4_text if block[1] == "rq4_interpretation" else "")
            continue
        if kind == "figref":
            number = block[1]
            fig_entry = manifest["figures"][number]
            png_name = Path(fig_entry["files"]["png"]).name
            doc.add_paragraph(f"See Figure {fig_entry['display_number']} ({fig_entry.get('caption', '')}) -- {png_name}.")
            continue

        # kind == "table"
        number = block[1]
        display_number = manifest["tables"].get(number, {}).get("display_number") or _display_number(number, args.numbering)
        caption = table_captions[number]
        if number == "4.8" and tables["4.8"] is None:
            missing_para = doc.add_paragraph()
            missing_run = missing_para.add_run(f"Table {display_number}")
            missing_run.bold = True
            missing_run.font.name = APA_FONT
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(caption)
            title_run.italic = True
            title_run.font.name = APA_FONT
            doc.add_paragraph("Unavailable: no ablation baseline checkpoint found. Run scripts/run_ablation.py first.")
            continue
        df = table_4_9 if number == "4.9" else tables[number][0]
        footnote = None
        if number == "4.8":
            footnote = f"Wilcoxon signed-rank test (baseline residuals > BIOPINN residuals): p = {tables['4.8'][1]['wilcoxon_p_value']:.3e}"
        _add_table_to_doc(doc, display_number, caption, df, footnote=footnote)

    doc.save(docx_path)
    print(f"  saved {docx_path}")

    # ----------------------------------------------------------------- #
    # Figure captions (APA style, external to the images -- see
    # _write_figure_captions)
    # ----------------------------------------------------------------- #
    figure_order = ["4.1", "4.2", "4.3", "4.4", "4.5", "4.6", "rq1b", "4.7", "4.8", "rq1c", "4.9", "4.10"]
    captions_path = _write_figure_captions(manifest["figures"], figure_order, args.numbering, output_dir)
    print(f"  saved {captions_path}")

    # ----------------------------------------------------------------- #
    # Manifest
    # ----------------------------------------------------------------- #
    manifest["research_questions"] = {
        "RQ1a": {"question": "Diffusion coefficient effect of nanoparticle diameter on drug penetration depth", "evidence": ["4.2", "4.4", "4.6"]},
        "RQ1b": {"question": "Diffusion coefficient effect of nanoparticle diameter on concentration distribution across tumor zones", "evidence": ["rq1b"]},
        "RQ1c": {"question": "Diffusion coefficient effect of nanoparticle diameter on sub-therapeutic regions", "evidence": ["rq1c", "4.6"]},
        "RQ2": {"question": "BIOPINN accuracy vs. the FDM reference on the held-out test set (RMSE, MAE, R², L2 relative error)", "evidence": ["4.3"]},
        "RQ3": {"question": "Predicted cell survivability and cytotoxicity during drug penetration and exposure", "evidence": ["4.5", "4.7", "4.8"]},
        "RQ4": {"question": "Nanoparticle diameter / dosing combination maximizing predicted tumor cell kill, and why", "evidence": ["4.7", "4.10"]},
    }
    manifest_path = output_dir / f"results_manifest_{args.numbering}.json"
    with open(manifest_path, "w", encoding="utf-8") as f:
        json.dump(_sanitize_for_json(manifest), f, indent=2)
    print(f"\nSaved {len(manifest['figures'])} figures (PNG+PDF), {len(manifest['tables'])} tables (CSV), {docx_path.name}, {captions_path.name}, and {manifest_path.name} to {output_dir}")

    print("\n--- Hypothesis summary ---")
    for key, h in hypotheses.items():
        print(f"  {key} ({h['name']}): {'SUPPORTED' if h['pass'] else 'NOT SUPPORTED'} -- {h['evidence']}")


if __name__ == "__main__":
    main()
